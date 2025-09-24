import os
import time
import tempfile
import telebot
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from pypdf import PdfReader

# Bot token
TOKEN = "8085614647:AAFg6oXkg0CdLeW2xoHMJ3lan53PGZjvIWE"
bot = telebot.TeleBot(TOKEN)

# Copyrights and contact info
BOT_RIGHTS = "🤍 تلجرام :- @altaee_z\n🌐موقعي : www.ali-Altaee.free.nf"

# Function to convert docx to PDF (text only)
def convert_docx_to_pdf_simple(input_path, output_path):
    doc = Document(input_path)
    pdf = SimpleDocTemplate(output_path)
    styles = getSampleStyleSheet()
    flow = []
    for para in doc.paragraphs:
        flow.append(Paragraph(para.text, styles["Normal"]))
        flow.append(Spacer(1, 12))
    pdf.build(flow)
    return os.path.exists(output_path)

@bot.message_handler(commands=['start', 'help'])
def send_welcome(message):
    bot.reply_to(message,
        "👋 أرسل ملف .docx وسأحوّله إلى PDF، مع عرض معلومات عنه.")

@bot.message_handler(content_types=['document'])
def handle_docs(message):
    doc = message.document
    filename = doc.file_name or "document.docx"
    chat_id = message.chat.id

    status = bot.reply_to(message, "⏳ جاري تحميل الملف...")
    start_time = time.time()

    try:
        file_info = bot.get_file(doc.file_id)
        downloaded = bot.download_file(file_info.file_path)

        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, filename)
            with open(input_path, 'wb') as f:
                f.write(downloaded)

            file_extension = os.path.splitext(filename)[1].lower()
            num_pages = None
            is_pdf_or_docx = False

            # Check file type to get page count
            if file_extension == '.docx':
                doc_obj = Document(input_path)
                num_pages = len(doc_obj.paragraphs)
                is_pdf_or_docx = True
            elif file_extension == '.pdf':
                reader = PdfReader(input_path)
                num_pages = len(reader.pages)
                is_pdf_or_docx = True
            
            # Get file size
            file_size_mb = os.path.getsize(input_path) / (1024*1024)
            
            # Calculate elapsed time
            elapsed = time.time() - start_time

            # Prepare the caption
            if is_pdf_or_docx:
                caption = (
                    f"✅ تم استلام الملف بنجاح!\n\n"
                    f"📑 تفاصيل الملف\n"
                    f"• الحجم: {file_size_mb:.2f} MB\n"
                    f"• عدد الصفحات: {num_pages}\n"
                    f"• الوقت المستغرق: {elapsed:.2f} ثانية\n\n"
                    f"{BOT_RIGHTS}"
                )
            else:
                caption = (
                    f"✅ تم استلام الملف بنجاح!\n\n"
                    f"📑 تفاصيل الملف\n"
                    f"• الحجم: {file_size_mb:.2f} MB\n"
                    f"• عدد الصفحات: غير متاح لهذا النوع من الملفات.\n"
                    f"• الوقت المستغرق: {elapsed:.2f} ثانية\n\n"
                    f"{BOT_RIGHTS}"
                )
            
            bot.edit_message_text("✅ تم إرسال التفاصيل...", chat_id, status.message_id)

            # Send the document with the new combined caption.
            with open(input_path, 'rb') as received_file:
                bot.send_document(chat_id, received_file, caption=caption)

    except Exception as e:
        bot.edit_message_text(f"❗ حدث خطأ: {e}", chat_id, status.message_id)

if __name__ == "__main__":
    print("🚀 Bot running...")
    bot.infinity_polling()
